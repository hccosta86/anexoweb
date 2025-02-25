from flask import Flask, render_template, request, redirect, url_for, send_file, flash
import sqlite3
import random
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

app = Flask(__name__)
app.secret_key = 'sua_chave_secreta_aqui'
app.config['UPLOAD_FOLDER'] = 'static/uploads'

# Extensões de arquivo permitidas para upload
ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png'}

# Função para verificar se a extensão do arquivo é permitida
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Função para criar o banco de dados
def criar_banco():
    conn = sqlite3.connect('servidores.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS servidores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            masp TEXT UNIQUE NOT NULL,
            nome TEXT NOT NULL,
            sexo TEXT NOT NULL,
            raca TEXT NOT NULL,
            foto TEXT NOT NULL,
            barba TEXT NOT NULL,
            careca TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

# Função para adicionar bordas às células da tabela
def set_cell_borders(cell, border_size=4):
    """Adiciona bordas a uma célula da tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top', 'left', 'bottom', 'right']:
        tag = f'w:{border_name}'
        element = OxmlElement(tag)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), str(border_size))
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        tcPr.append(element)

# Rota principal
@app.route('/')
def index():
    return render_template('index.html')

# Rota para cadastrar servidores
@app.route('/cadastrar', methods=['GET', 'POST'])
def cadastrar():
    if request.method == 'POST':
        masp = request.form.get('masp', '').strip()
        nome = request.form.get('nome', '').strip()
        sexo = request.form.get('sexo', '').strip()
        raca = request.form.get('raca', '').strip()
        barba = request.form.get('barba', 'Não')
        careca = request.form.get('careca', 'Não')
        foto = request.files.get('foto')

        # Extrai nome e MASP do nome da foto (se aplicável)
        if foto and allowed_file(foto.filename):
            nome_arquivo = os.path.splitext(foto.filename)[0]  # Remove a extensão
            partes = nome_arquivo.split('_')  # Divide o nome do arquivo
            if len(partes) >= 2:
                nome = ' '.join([parte.capitalize() for parte in partes[:-1]])  # Nome completo
                masp = partes[-1]  # Última parte é o MASP

        # Validação dos campos obrigatórios
        if not masp or not nome or not sexo or not raca:
            flash('Todos os campos são obrigatórios.', 'error')
            return redirect(url_for('cadastrar'))

        # Validação do arquivo de foto
        if not foto or not allowed_file(foto.filename):
            flash('Por favor, envie uma foto válida (JPG, JPEG ou PNG).', 'error')
            return redirect(url_for('cadastrar'))

        # Salvar a foto no diretório de uploads
        try:
            foto_path = os.path.join(app.config['UPLOAD_FOLDER'], foto.filename)
            foto.save(foto_path)

            # Salvar o caminho completo da foto no banco de dados
            foto_path_db = foto_path  # Caminho completo
        except Exception as e:
            flash(f'Erro ao salvar a foto: {e}', 'error')
            return redirect(url_for('cadastrar'))

        # Inserir os dados no banco de dados
        conn = sqlite3.connect('servidores.db')
        cursor = conn.cursor()
        try:
            cursor.execute('''
                INSERT INTO servidores (masp, nome, sexo, raca, foto, barba, careca)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (masp, nome, sexo, raca, foto_path_db, barba, careca))
            conn.commit()
            flash('Servidor cadastrado com sucesso!', 'success')
        except sqlite3.IntegrityError:
            flash('MASP já cadastrado.', 'error')
        except sqlite3.Error as e:
            flash(f'Erro ao cadastrar servidor: {e}', 'error')
        finally:
            conn.close()

        return redirect(url_for('listar'))
    return render_template('cadastro.html')

# Rota para listar servidores
@app.route('/listar')
def listar():
    conn = sqlite3.connect('servidores.db')
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM servidores')
    servidores = cursor.fetchall()
    conn.close()

    # Ajusta o caminho das fotos para o template
    servidores_ajustados = []
    for servidor in servidores:
        servidor_lista = list(servidor)
        if servidor_lista[5]:  # Verifica se há um caminho de foto
            servidor_lista[5] = servidor_lista[5].replace('static/', '')  # Remove o prefixo 'static/'
        servidores_ajustados.append(servidor_lista)

    return render_template('listagem.html', servidores=servidores_ajustados)

# Rota para excluir servidor
@app.route('/excluir/<int:id>')
def excluir(id):
    conn = sqlite3.connect('servidores.db')
    cursor = conn.cursor()
    try:
        cursor.execute('DELETE FROM servidores WHERE id = ?', (id,))
        conn.commit()
        flash('Servidor excluído com sucesso!', 'success')
    except sqlite3.Error as e:
        flash(f'Erro ao excluir servidor: {e}', 'error')
    finally:
        conn.close()
    return redirect(url_for('listar'))

# Rota para seleção de servidores (obrigatórios e aleatórios)
@app.route('/anexo', methods=['GET', 'POST'])
def anexo():
    if request.method == 'POST':
        servidores_obrigatorios = request.form.getlist('obrigatorios')
        quantidade_aleatorios = int(request.form.get('quantidade_aleatorios', 0))
        filtro_sexo = request.form.get('filtro_sexo', 'Todos')
        filtro_raca = request.form.get('filtro_raca', 'Todos')
        filtro_barba = request.form.get('filtro_barba', 'Todos')
        filtro_careca = request.form.get('filtro_careca', 'Todos')

        conn = sqlite3.connect('servidores.db')
        cursor = conn.cursor()

        # Seleciona servidores obrigatórios
        cursor.execute('SELECT * FROM servidores WHERE id IN ({})'.format(','.join('?' for _ in servidores_obrigatorios)), servidores_obrigatorios)
        servidores_obrigatorios = cursor.fetchall()

        # Filtra servidores disponíveis para aleatoriedade
        query_filtro = 'SELECT * FROM servidores WHERE id NOT IN ({})'.format(','.join('?' for _ in servidores_obrigatorios))
        params_filtro = [servidor[0] for servidor in servidores_obrigatorios]

        # Aplica filtros
        if filtro_sexo != 'Todos':
            query_filtro += ' AND sexo = ?'
            params_filtro.append(filtro_sexo)
        if filtro_raca != 'Todos':
            query_filtro += ' AND raca = ?'
            params_filtro.append(filtro_raca)
        if filtro_barba != 'Todos':
            query_filtro += ' AND barba = ?'
            params_filtro.append(filtro_barba)
        if filtro_careca != 'Todos':
            query_filtro += ' AND careca = ?'
            params_filtro.append(filtro_careca)

        cursor.execute(query_filtro, params_filtro)
        servidores_disponiveis = cursor.fetchall()

        # Seleciona servidores aleatórios
        servidores_aleatorios = random.sample(servidores_disponiveis, min(quantidade_aleatorios, len(servidores_disponiveis)))

        # Combina as listas
        lista_final = servidores_obrigatorios + servidores_aleatorios
        conn.close()

        # Ajusta o caminho das fotos para o template
        lista_final_ajustada = []
        for servidor in lista_final:
            servidor_lista = list(servidor)
            if servidor_lista[5]:  # Verifica se há um caminho de foto
                servidor_lista[5] = servidor_lista[5].replace('static/', '')  # Remove o prefixo 'static/'
            lista_final_ajustada.append(servidor_lista)

        return render_template('lista_final.html', lista_final=lista_final_ajustada)

    conn = sqlite3.connect('servidores.db')
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM servidores')
    servidores = cursor.fetchall()
    conn.close()

    # Ajusta o caminho das fotos para o template
    servidores_ajustados = []
    for servidor in servidores:
        servidor_lista = list(servidor)
        if servidor_lista[5]:  # Verifica se há um caminho de foto
            servidor_lista[5] = servidor_lista[5].replace('static/', '')  # Remove o prefixo 'static/'
        servidores_ajustados.append(servidor_lista)

    return render_template('anexo.html', servidores=servidores_ajustados)

# Rota para gerar o anexo fotográfico
@app.route('/gerar_anexo', methods=['POST'])
def gerar_anexo():
    servidores_selecionados = request.form.getlist('servidores')
    conn = sqlite3.connect('servidores.db')
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM servidores WHERE id IN ({})'.format(','.join('?' for _ in servidores_selecionados)), servidores_selecionados)
    servidores = cursor.fetchall()
    conn.close()

    # Gerar DOCX
    doc = Document()
    doc.add_heading('Anexo Fotográfico', 0)

    # Adiciona as fotos lado a lado (4 por linha)
    fotos_por_linha = 4
    num_servidores = len(servidores)
    num_linhas = (num_servidores + fotos_por_linha - 1) // fotos_por_linha

    tabela_fotos = doc.add_table(rows=num_linhas, cols=fotos_por_linha)
    tabela_fotos.autofit = False

    # Define o tamanho das células
    largura_celula = Inches(1.5)
    for row in tabela_fotos.rows:
        for cell in row.cells:
            cell.width = largura_celula

    # Adiciona as fotos e os números de referência na tabela
    for i, servidor in enumerate(servidores):
        linha = i // fotos_por_linha
        coluna = i % fotos_por_linha

        # Adiciona o número de referência acima da foto
        cell = tabela_fotos.cell(linha, coluna)
        cell.text = f"[{i + 1}]"
        cell.paragraphs[0].alignment = 1  # Centraliza o texto

        # Adiciona a foto
        if servidor[5]:  # Verifica se o caminho da foto existe
            foto_path = servidor[5]  # Caminho completo da foto
            if os.path.exists(foto_path):
                try:
                    paragraph = cell.add_paragraph()
                    paragraph.alignment = 1  # Centraliza a foto
                    run = paragraph.add_run()
                    run.add_picture(foto_path, width=Inches(1.5))
                except Exception as e:
                    print(f"Erro ao adicionar a foto: {e}")
                    flash(f"Erro ao adicionar a foto: {e}", "error")
            else:
                print(f"Arquivo de foto não encontrado: {foto_path}")
                flash(f"Arquivo de foto não encontrado: {foto_path}", "error")
        else:
            print(f"Servidor {servidor[2]} não tem foto cadastrada.")
            flash(f"Servidor {servidor[2]} não tem foto cadastrada.", "error")

    # Adiciona uma nova página para a lista de referência
    doc.add_page_break()
    doc.add_heading('Lista de Referência', 0)

    # Adiciona a lista de referência em uma tabela
    tabela_referencia = doc.add_table(rows=1, cols=3)
    tabela_referencia.autofit = False

    # Define o cabeçalho da tabela com formatação
    cabecalho = tabela_referencia.rows[0].cells
    cabecalho[0].text = "Número"
    cabecalho[1].text = "Nome"
    cabecalho[2].text = "MASP"

    # Formatação do cabeçalho (negrito e fonte maior)
    for cell in cabecalho:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)

    # Adiciona os dados dos servidores na tabela
    for i, servidor in enumerate(servidores, start=1):
        row = tabela_referencia.add_row().cells
        row[0].text = str(i)
        row[1].text = servidor[2]  # Nome
        row[2].text = servidor[1]  # MASP

    # Adiciona bordas às células da tabela de referência
    for row in tabela_referencia.rows:
        for cell in row.cells:
            set_cell_borders(cell)

    # Salva o documento em um arquivo temporário para depuração
    temp_file_path = "temp_anexo_fotografico.docx"
    doc.save(temp_file_path)
    print(f"Documento salvo temporariamente em: {temp_file_path}")  # Depuração

    # Retorna o arquivo para o usuário
    return send_file(temp_file_path, as_attachment=True, download_name='anexo_fotografico.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    criar_banco()
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    app.run(debug=True)